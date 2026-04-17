import io
import re
from fastapi import APIRouter, HTTPException, Body
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter()


class ExportWordRequest(BaseModel):
    html_content: str
    filename: str = "contrato_editado"


def _html_to_docx(html: str) -> io.BytesIO:
    """
    Converte HTML do Tiptap em um .docx de nível profissional usando python-docx.
    Suporta: parágrafos, headings (h1-h3), negrito, itálico, listas.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import html as html_lib
    from html.parser import HTMLParser

    doc = Document()

    # --- Configuração de Página A4 com margens padrão jurídico ---
    # NOTA: Cm() converte centímetros para EMUs corretamente (1cm = 360.000 EMUs)
    section = doc.sections[0]
    section.page_width   = Cm(21.0)    # A4: 21 cm
    section.page_height  = Cm(29.7)    # A4: 29,7 cm
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

    # --- Estilo Normal base ---
    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(18)  # 1.5 lines

    class TiptapParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.current_para = None
            self.current_run_bold = False
            self.current_run_italic = False
            self.in_list_item = False
            self.list_level = 0

        def _ensure_para(self):
            if self.current_para is None:
                self.current_para = doc.add_paragraph()
                self.current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag == 'p':
                self.current_para = doc.add_paragraph()
                self.current_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self.in_list_item = False
            elif tag == 'h1':
                self.current_para = doc.add_heading('', level=1)
            elif tag == 'h2':
                self.current_para = doc.add_heading('', level=2)
            elif tag == 'h3':
                self.current_para = doc.add_heading('', level=3)
            elif tag in ('strong', 'b'):
                self.current_run_bold = True
            elif tag in ('em', 'i'):
                self.current_run_italic = True
            elif tag == 'li':
                self.current_para = doc.add_paragraph(style='List Bullet')
                self.in_list_item = True
            elif tag in ('ul', 'ol'):
                self.list_level += 1

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in ('p', 'h1', 'h2', 'h3', 'li'):
                self.current_para = None
                self.in_list_item = False
            elif tag in ('strong', 'b'):
                self.current_run_bold = False
            elif tag in ('em', 'i'):
                self.current_run_italic = False
            elif tag in ('ul', 'ol'):
                self.list_level = max(0, self.list_level - 1)

        def handle_data(self, data):
            import html as html_lib
            text = html_lib.unescape(data)
            # Ignora nós que são apenas whitespace/quebras entre tags (ex: "\n  " entre <p> e <strong>)
            if not text or not text.strip():
                return
            # Remove quebras de linha internas — no Word a própria página faz o wrap
            text = text.replace('\n', ' ').replace('\r', '')
            self._ensure_para()
            run = self.current_para.add_run(text)
            run.bold   = self.current_run_bold
            run.italic = self.current_run_italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    parser = TiptapParser()
    # Remove atributos de estilo inline e data-attrs do Tiptap antes de parsear
    clean_html = re.sub(r'\s*(style|data-[a-z\-]+)="[^"]*"', '', html)
    # Remove tags que o Word não precisa processar
    clean_html = re.sub(r'<br\s*/?>', ' ', clean_html)
    parser.feed(clean_html)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.post("/word", response_class=StreamingResponse)
async def export_to_word(payload: ExportWordRequest):
    """
    Recebe o HTML atual do editor Tiptap e gera um .docx profissional.
    Retorna StreamingResponse para download imediato.
    """
    try:
        docx_buffer = _html_to_docx(payload.html_content)
        safe_name = re.sub(r'[^\w\-]', '_', payload.filename) or "contrato_editado"

        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha na geração do Word: {str(e)}")
